from __future__ import annotations

import hashlib
import os
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

# Avoid optional TensorFlow import path in transformers/sentence-transformers.
os.environ.setdefault("USE_TF", "0")
os.environ.setdefault("TRANSFORMERS_NO_TF", "1")

from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from PyPDF2 import PdfReader


DEFAULT_GROQ_MODEL = os.getenv("GROQ_MODEL", "llama-3.1-8b-instant")


@dataclass
class BackendError(Exception):
    code: str
    message: str
    details: Optional[Any] = None

    def to_dict(self) -> Dict[str, Any]:
        d: Dict[str, Any] = {"code": self.code, "message": self.message}
        if self.details is not None:
            d["details"] = self.details
        return d

    def __str__(self) -> str:
        if self.details is None:
            return f"{self.code}: {self.message}"
        return f"{self.code}: {self.message} | details={self.details}"


def _file_ext(filename: str) -> str:
    parts = filename.rsplit(".", 1)
    return parts[-1].lower() if len(parts) == 2 else ""


def _ensure_api_key() -> None:
    if os.getenv("GROQ_API_KEY"):
        return

    # Fallback: load from local Streamlit secrets when running outside Streamlit process
    # (e.g., FastAPI API mode).
    try:
        import tomllib  # Python 3.11+
    except Exception:
        tomllib = None  # type: ignore[assignment]

    if tomllib is not None:
        try:
            secrets_path = Path(".streamlit/secrets.toml")
            if secrets_path.exists():
                data = tomllib.loads(secrets_path.read_text(encoding="utf-8"))
                key = str(data.get("GROQ_API_KEY") or "").strip()
                if key:
                    os.environ["GROQ_API_KEY"] = key
                    return
        except Exception:
            pass

    if not os.getenv("GROQ_API_KEY"):
        raise BackendError(
            code="MISSING_GROQ_API_KEY",
            message="Groq API key bulunamadı. GROQ_API_KEY ortam değişkenini ayarlayın veya .streamlit/secrets.toml içine ekleyin.",
        )


# ── In-memory FAISS store (no disk I/O, no SQLite) ──────────────────────────
_EMBEDDINGS: Optional[HuggingFaceEmbeddings] = None
_VECTORSTORE: Optional[FAISS] = None
_INDEX_ID: Optional[str] = None
_VECTOR_DB_DIR = Path(os.getenv("VECTOR_DB_PATH", ".rag_faiss")).resolve()
_LATEST_INDEX_FILE = _VECTOR_DB_DIR / "latest_index.txt"


def _get_embeddings() -> HuggingFaceEmbeddings:
    global _EMBEDDINGS
    if _EMBEDDINGS is None:
        _EMBEDDINGS = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
    return _EMBEDDINGS


def _reset_vectorstore() -> None:
    global _VECTORSTORE, _INDEX_ID
    _VECTORSTORE = None
    _INDEX_ID = None


def _index_dir(index_id: str) -> Path:
    return _VECTOR_DB_DIR / index_id


def _write_latest_index_id(index_id: str) -> None:
    _VECTOR_DB_DIR.mkdir(parents=True, exist_ok=True)
    _LATEST_INDEX_FILE.write_text(index_id, encoding="utf-8")


def _read_latest_index_id() -> Optional[str]:
    if not _LATEST_INDEX_FILE.exists():
        return None
    value = _LATEST_INDEX_FILE.read_text(encoding="utf-8").strip()
    return value or None


def _persist_vectorstore(vs: FAISS, index_id: str) -> Path:
    out_dir = _index_dir(index_id)
    out_dir.mkdir(parents=True, exist_ok=True)
    vs.save_local(str(out_dir))
    _write_latest_index_id(index_id)
    return out_dir


def _load_vectorstore(index_id: str) -> FAISS:
    in_dir = _index_dir(index_id)
    if not in_dir.exists():
        raise BackendError(
            "INDEX_NOT_FOUND",
            "İstenen vektör indeksi bulunamadı.",
            {"index_id": index_id, "path": str(in_dir)},
        )
    return FAISS.load_local(
        str(in_dir),
        _get_embeddings(),
        allow_dangerous_deserialization=True,
    )


def _compute_index_id(docs: List[Document]) -> str:
    digest = hashlib.sha1()
    digest.update(f"doc_count:{len(docs)}|".encode("utf-8"))
    for d in docs[:128]:
        meta = d.metadata or {}
        source = str(meta.get("source", ""))
        chunk = str(meta.get("chunk", ""))
        text = (d.page_content or "")[:300]
        digest.update(f"{source}:{chunk}:".encode("utf-8", errors="ignore"))
        digest.update(text.encode("utf-8", errors="ignore"))
    return digest.hexdigest()


def _build_vectorstore(docs: List[Document]) -> Tuple[FAISS, str]:
    embeddings = _get_embeddings()
    # Build in-memory first, then persist on disk.
    vs = FAISS.from_documents(documents=docs, embedding=embeddings)
    index_id = _compute_index_id(docs)
    return vs, index_id


# ── Text extraction helpers ──────────────────────────────────────────────────

def _extract_text_pdf(file_obj) -> Tuple[str, List[Dict[str, Any]]]:
    reader = PdfReader(file_obj)
    pages_meta: List[Dict[str, Any]] = []
    texts: List[str] = []
    for i, page in enumerate(reader.pages):
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        texts.append(t)
        pages_meta.append({"page": i + 1, "chars": len(t)})
    return "\n\n".join(texts).strip(), pages_meta


def _extract_text_docx(file_obj) -> str:
    doc = DocxDocument(file_obj)
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    return "\n".join(parts).strip()


def _extract_text_txt(file_obj) -> str:
    raw = file_obj.read()
    if isinstance(raw, str):
        return raw.strip()
    for enc in ("utf-8", "utf-8-sig", "cp1254", "latin-1"):
        try:
            return raw.decode(enc).strip()
        except Exception:
            continue
    return raw.decode("utf-8", errors="ignore").strip()


def _extract_text_doc(file_obj) -> str:
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise BackendError(
            "DOC_UNSUPPORTED",
            "DOC desteği için LibreOffice gerekli. Dosyayı DOCX olarak yükleyin.",
            {"hint": "LibreOffice kurun veya DOCX yükleyin"},
        )

    data = file_obj.read()
    if isinstance(data, str):
        data = data.encode("utf-8", errors="ignore")

    with tempfile.TemporaryDirectory() as td:
        td_path = Path(td)
        in_path = td_path / "input.doc"
        out_dir = td_path / "out"
        out_dir.mkdir(parents=True, exist_ok=True)
        in_path.write_bytes(data)

        cmd = [
            soffice, "--headless", "--nologo", "--nolockcheck",
            "--nodefault", "--norestore", "--convert-to", "docx",
            "--outdir", str(out_dir), str(in_path),
        ]
        try:
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        except Exception as e:
            raise BackendError("DOC_CONVERT_FAILED", "DOC → DOCX dönüşümü başarısız.", str(e))

        out_path = out_dir / "input.docx"
        if not out_path.exists():
            candidates = list(out_dir.glob("*.docx"))
            if not candidates:
                raise BackendError("DOC_CONVERT_FAILED", "DOC → DOCX dönüşümü çıktı üretmedi.")
            out_path = candidates[0]

        with out_path.open("rb") as f:
            return _extract_text_docx(f)


def _documents_from_uploaded_files(
    uploaded_files: Iterable,
) -> Tuple[List[Document], List[Dict[str, Any]]]:
    extracted: List[Document] = []
    errors: List[Dict[str, Any]] = []

    for uf in uploaded_files:
        name = getattr(uf, "name", "dosya")
        ext = _file_ext(name)

        try:
            if ext == "pdf":
                text, pages_meta = _extract_text_pdf(uf)
                if not text:
                    raise BackendError("DOC_EMPTY", "PDF'den metin çıkarılamadı.", {"filename": name, "pages": pages_meta})
                extracted.append(Document(page_content=text, metadata={"source": name, "type": "pdf"}))

            elif ext == "docx":
                text = _extract_text_docx(uf)
                if not text:
                    raise BackendError("DOC_EMPTY", "DOCX'ten metin çıkarılamadı.", {"filename": name})
                extracted.append(Document(page_content=text, metadata={"source": name, "type": "docx"}))

            elif ext == "doc":
                text = _extract_text_doc(uf)
                if not text:
                    raise BackendError("DOC_EMPTY", "DOC'tan metin çıkarılamadı.", {"filename": name})
                extracted.append(Document(page_content=text, metadata={"source": name, "type": "doc"}))

            elif ext == "txt":
                text = _extract_text_txt(uf)
                if not text:
                    raise BackendError("DOC_EMPTY", "TXT dosyası boş.", {"filename": name})
                extracted.append(Document(page_content=text, metadata={"source": name, "type": "txt"}))

            else:
                errors.append({"filename": name, "reason": "unsupported_file_type"})

        except BackendError as e:
            errors.append({"filename": name, "reason": e.code, "message": e.message, "details": e.details})
        except Exception as e:
            errors.append({"filename": name, "reason": "extract_failed", "message": str(e)})

    return extracted, errors


# ── Public API ───────────────────────────────────────────────────────────────

def process_documents(uploaded_files: List) -> Dict[str, Any]:
    """
    Streamlit UploadedFile listesi alır.
    Dökümanları okur → chunk'lar → FAISS in-memory index kurar.
    """
    global _VECTORSTORE, _INDEX_ID

    if not uploaded_files:
        return {"ok": False, "error": {"code": "NO_FILES", "message": "İşlenecek dosya bulunamadı."}}

    raw_docs, errors = _documents_from_uploaded_files(uploaded_files)
    if not raw_docs:
        return {
            "ok": False,
            "error": {
                "code": "NO_EXTRACTED_TEXT",
                "message": "Hiçbir dosyadan metin çıkarılamadı.",
                "details": errors,
            },
        }

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=150,
        separators=["\n\n", "\n", " ", ""],
    )

    chunked: List[Document] = []
    for d in raw_docs:
        chunks = splitter.split_text(d.page_content)
        for idx, chunk in enumerate(chunks):
            chunked.append(
                Document(
                    page_content=chunk,
                    metadata={**(d.metadata or {}), "chunk": idx},
                )
            )

    try:
        vs, index_id = _build_vectorstore(chunked)
    except Exception as e:
        return {
            "ok": False,
            "error": {"code": "VECTORSTORE_FAILED", "message": f"Vektör indeksi oluşturulamadı: {e}"},
        }

    _VECTORSTORE = vs
    _INDEX_ID = index_id

    response: Dict[str, Any] = {
        "ok": True,
        "index_id": index_id,
        "processed_count": len(raw_docs),
        "stats": {
            "chunks": len(chunked),
            "embedding_model": "sentence-transformers/all-MiniLM-L6-v2",
            "vector_store": "FAISS (disk + memory)",
        },
        "errors": errors,
    }

    try:
        saved_path = _persist_vectorstore(vs, index_id)
        response["storage"] = {
            "index_id": index_id,
            "path": str(saved_path),
            "persistent": True,
        }
    except Exception as e:
        response["storage"] = {
            "index_id": index_id,
            "persistent": False,
            "warning": f"İndeks diske kaydedilemedi: {e}",
        }

    return response


def _require_vectorstore(index_id: Optional[str] = None) -> Tuple[FAISS, str]:
    global _VECTORSTORE, _INDEX_ID

    if _VECTORSTORE is not None and _INDEX_ID is not None and (index_id is None or index_id == _INDEX_ID):
        return _VECTORSTORE, _INDEX_ID

    target_index_id = index_id or _read_latest_index_id()
    if target_index_id:
        try:
            loaded = _load_vectorstore(target_index_id)
            _VECTORSTORE = loaded
            _INDEX_ID = target_index_id
            return loaded, target_index_id
        except BackendError:
            raise
        except Exception as e:
            raise BackendError("INDEX_LOAD_FAILED", "Vektör indeksi yüklenemedi.", str(e))

    raise BackendError(
        "INDEX_NOT_FOUND",
        "Vektör indeksi bulunamadı. Lütfen dökümanları tekrar işleyin.",
    )


def summarize_documents(
    *,
    index_id: Optional[str] = None,
    max_chars: int = 12000,
) -> Dict[str, Any]:
    """Yüklenen dökümanların tamamından kısa, kaynaklı bir özet üretir."""
    _ensure_api_key()

    vs, active_index_id = _require_vectorstore(index_id)
    model = os.getenv("GROQ_MODEL", DEFAULT_GROQ_MODEL)

    doc_dict = getattr(getattr(vs, "docstore", None), "_dict", {}) or {}
    docs = list(doc_dict.values())
    if not docs:
        return {
            "ok": True,
            "index_id": active_index_id,
            "summary": "Özet oluşturmak için işlenmiş döküman bulunamadı.",
            "sources": [],
            "meta": {"model": model},
        }

    context_parts: List[str] = []
    used_chars = 0
    sources: List[Dict[str, Any]] = []

    for i, d in enumerate(docs):
        text = (getattr(d, "page_content", "") or "").strip()
        if not text:
            continue

        meta = getattr(d, "metadata", {}) or {}
        filename = meta.get("source", "doküman")
        chunk_id = meta.get("chunk")

        snippet = text.replace("\n", " ")
        snippet = snippet[:220] + ("…" if len(snippet) > 220 else "")
        sources.append({"filename": filename, "chunk_id": chunk_id, "snippet": snippet})

        remaining = max_chars - used_chars
        if remaining <= 0:
            break
        piece = text[:remaining]
        context_parts.append(f"[Kaynak: {filename} | chunk={chunk_id}]\n{piece}")
        used_chars += len(piece)

        if i >= 59:
            break

    context = "\n\n---\n\n".join(context_parts).strip()
    if not context:
        return {
            "ok": True,
            "index_id": active_index_id,
            "summary": "Dökümanlardan anlamlı içerik çıkarılamadı.",
            "sources": [],
            "meta": {"model": model},
        }

    prompt = ChatPromptTemplate.from_messages([
        (
            "system",
            "Sen doküman özetleme asistanısın. Sadece verilen bağlama göre Türkçe, kısa, net ve maddeli özet çıkar. "
            "Bağlam dışı bilgi üretme.",
        ),
        (
            "human",
            "Bağlam (context):\n{context}\n\n"
            "İstenen çıktı formatı:\n"
            "1) Genel özet (2-4 cümle)\n"
            "2) Öne çıkan bulgular (3-6 madde)\n"
            "3) Kritik rakam/tarih/isimler (varsa)\n"
            "4) Belirsiz kalan noktalar\n\n"
            "Cevap:",
        ),
    ])

    llm = ChatGroq(model=model, temperature=0.2, max_tokens=900)
    try:
        msg = llm.invoke(prompt.format_messages(context=context))
        summary = (getattr(msg, "content", None) or "").strip()
    except Exception as e:
        raise BackendError("LLM_FAILED", "Groq özetleme çağrısı başarısız.", str(e))

    if not summary:
        summary = "Özet üretilemedi (boş yanıt)."

    return {
        "ok": True,
        "index_id": active_index_id,
        "summary": summary,
        "sources": sources[:12],
        "meta": {"model": model, "max_chars": max_chars},
    }


def get_ai_response(
    user_query: str,
    *,
    index_id: Optional[str] = None,
    chat_history: Optional[List[Dict[str, str]]] = None,
    top_k: int = 4,
) -> Dict[str, Any]:
    """Retrieval + Groq LLM ile döküman tabanlı cevap üretir."""
    _ensure_api_key()

    if not user_query or not user_query.strip():
        return {"ok": False, "error": {"code": "EMPTY_QUERY", "message": "Soru boş olamaz."}}

    vs, active_index_id = _require_vectorstore(index_id)

    try:
        results = vs.similarity_search_with_score(user_query, k=top_k)
    except Exception as e:
        raise BackendError("RETRIEVAL_FAILED", "Retrieval işlemi başarısız.", str(e))

    scored_sources: List[Dict[str, Any]] = []
    context_parts: List[str] = []

    for d, score in results:
        meta = d.metadata or {}
        filename = meta.get("source", "doküman")
        chunk_id = meta.get("chunk")
        snippet = (d.page_content or "").strip().replace("\n", " ")
        snippet = snippet[:260] + ("…" if len(snippet) > 260 else "")

        scored_sources.append({
            "filename": filename,
            "chunk_id": chunk_id,
            "score": float(score),
            "snippet": snippet,
        })
        context_parts.append(f"[Kaynak: {filename} | chunk={chunk_id}]\n{d.page_content}")

    context = "\n\n---\n\n".join(context_parts).strip()
    if not context:
        return {
            "ok": True,
            "index_id": active_index_id,
            "answer": "Yüklediğiniz dökümanlarda bu soruya dair yeterli bilgi bulamadım.",
            "sources": [],
        }

    system = (
        "Sen bir doküman tabanlı asistanısın. "
        "Kullanıcı sorularını SADECE sağlanan bağlam (context) ile yanıtla. "
        "Bağlamda yoksa bunu açıkça söyle ve varsayım yapma. "
        "Cevabı Türkçe ve net yaz."
    )

    prompt = ChatPromptTemplate.from_messages([
        ("system", system),
        ("human", "Bağlam (context):\n{context}\n\nSoru:\n{question}\n\nCevap:"),
    ])

    model = os.getenv("GROQ_MODEL", DEFAULT_GROQ_MODEL)
    llm = ChatGroq(model=model, temperature=0.2, max_tokens=800)

    try:
        msg = llm.invoke(prompt.format_messages(context=context, question=user_query))
        answer = (getattr(msg, "content", None) or "").strip()
    except Exception as e:
        raise BackendError("LLM_FAILED", "Groq LLM çağrısı başarısız.", str(e))

    if not answer:
        answer = "Yanıt üretirken bir sorun oluştu (boş yanıt)."

    return {
        "ok": True,
        "index_id": active_index_id,
        "answer": answer,
        "sources": scored_sources,
        "meta": {"top_k": top_k, "model": model},
    }


def stream_ai_response(
    user_query: str,
    *,
    index_id: Optional[str] = None,
    chat_history: Optional[List[Dict[str, str]]] = None,
    top_k: int = 4,
):
    """Token streaming generator için FAISS tabanlı retrieval."""
    _ensure_api_key()

    if not user_query or not user_query.strip():
        raise BackendError("EMPTY_QUERY", "Soru boş olamaz.")

    vs, active_index_id = _require_vectorstore(index_id)

    try:
        results = vs.similarity_search_with_score(user_query, k=top_k)
    except Exception as e:
        raise BackendError("RETRIEVAL_FAILED", "Retrieval işlemi başarısız.", str(e))

    scored_sources: List[Dict[str, Any]] = []
    context_parts: List[str] = []

    for d, score in results:
        meta = d.metadata or {}
        filename = meta.get("source", "doküman")
        chunk_id = meta.get("chunk")
        snippet = (d.page_content or "").strip().replace("\n", " ")
        snippet = snippet[:260] + ("…" if len(snippet) > 260 else "")

        scored_sources.append({"filename": filename, "chunk_id": chunk_id, "score": float(score), "snippet": snippet})
        context_parts.append(f"[Kaynak: {filename} | chunk={chunk_id}]\n{d.page_content}")

    context = "\n\n---\n\n".join(context_parts).strip()
    if not context:
        yield ("token", "Yüklediğiniz dökümanlarda bu soruya dair yeterli bilgi bulamadım.")
        yield ("sources", [])
        return

    system = (
        "Sen bir doküman tabanlı asistanısın. "
        "Kullanıcı sorularını SADECE sağlanan bağlam (context) ile yanıtla. "
        "Bağlamda yoksa bunu açıkça söyle ve varsayım yapma. "
        "Cevabı Türkçe ve net yaz."
    )
    prompt = ChatPromptTemplate.from_messages([
        ("system", system),
        ("human", "Bağlam (context):\n{context}\n\nSoru:\n{question}\n\nCevap:"),
    ])

    model = os.getenv("GROQ_MODEL", DEFAULT_GROQ_MODEL)
    llm = ChatGroq(model=model, temperature=0.2, max_tokens=800)

    try:
        for chunk in llm.stream(prompt.format_messages(context=context, question=user_query)):
            text = (getattr(chunk, "content", None) or "")
            if text:
                yield ("token", text)
    except Exception as e:
        raise BackendError("LLM_FAILED", "Groq LLM streaming başarısız.", str(e))

    yield ("sources", scored_sources)